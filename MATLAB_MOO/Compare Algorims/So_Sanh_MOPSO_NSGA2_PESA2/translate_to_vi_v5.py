# -*- coding: utf-8 -*-
"""
translate_to_vi_v5.py
----------------------
Dich toan bo bai bao (dua tren "ICERA. V4.docx", ban da chinh van phong)
sang tieng Viet, xuat ra "ICERA. V5TV.docx".

Nguyen tac:
- Dich: tieu de, tom tat, tu khoa, tat ca tieu de muc, doan van xuoi,
  chu thich bang/hinh, nhan (label) trong bang so lieu.
- GIU NGUYEN, KHONG dich/dong vao: cong thuc toan (m:oMath), so lieu/ket
  qua trong bang, ky hieu bien (x1..x8, rho, Ai, Li, Ne, Fk, fk, gj...),
  ten thuat toan (MOPSO/NSGA-II/PESA-II), va toan bo danh muc tai lieu
  tham khao (giu nguyen ngon ngu goc cua cac cong bo duoc trich dan).
- Voi cac doan co cong thuc chen giua cau (paragraph 11,13,14,15), chi
  thay noi dung o cac "doan van ban" xen giua, khong dong vao cac phan
  tu <m:oMath> de tranh lam mat/hong cong thuc.
"""

import docx
from docx.oxml.ns import qn

SRC = "ICERA. V4.docx"
OUT = "ICERA. V5TV.docx"


def local(tag):
    return tag.split("}")[-1] if "}" in tag else tag


# --------------------------------------------------------------------
# 1) Thay toan bo noi dung 1 doan van (khong co cong thuc chen giua)
# --------------------------------------------------------------------

def set_paragraph_text(paragraph, new_text, bold_prefix=None, font_pt=None):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    if bold_prefix:
        r1 = paragraph.add_run(bold_prefix)
        r1.bold = True
        if font_pt:
            r1.font.size = docx.shared.Pt(font_pt)
        r2 = paragraph.add_run(new_text)
        if font_pt:
            r2.font.size = docx.shared.Pt(font_pt)
    else:
        r = paragraph.add_run(new_text)
        if font_pt:
            r.font.size = docx.shared.Pt(font_pt)


# --------------------------------------------------------------------
# 2) Thay cac "doan van ban" xen giua cong thuc <m:oMath>, GIU NGUYEN
#    cac phan tu oMath tai dung vi tri cua chung trong doan van.
# --------------------------------------------------------------------

def replace_text_segments(p_element, new_texts):
    groups = []
    current = []
    for child in list(p_element):
        if local(child.tag) == "r":
            current.append(child)
        else:
            if current:
                groups.append(current)
                current = []
    if current:
        groups.append(current)

    if len(groups) != len(new_texts):
        raise ValueError(
            f"So nhom van ban ({len(groups)}) khong khop so doan thay the "
            f"({len(new_texts)}) - kiem tra lai cau truc doan van."
        )

    for group, text in zip(groups, new_texts):
        first = group[0]
        t_elements = [c for c in first if local(c.tag) == "t"]
        if t_elements:
            t_elements[0].text = text
            t_elements[0].set(qn("xml:space"), "preserve")
            for extra in t_elements[1:]:
                first.remove(extra)
        for r in group[1:]:
            r.getparent().remove(r)


# --------------------------------------------------------------------
# 3) Noi dung dich - thay toan bo doan van (khong co cong thuc chen giua)
#    { chi_so_doan_van : (bold_prefix_hoac_None, van_ban_moi, co_ep_font_10pt) }
# --------------------------------------------------------------------

WHOLE = {
    0: (None, "Tối ưu hóa đa mục tiêu khung thép tiết diện chữ I theo TCVN 5575:2024: "
              "Nghiên cứu so sánh các thuật toán metaheuristic", False),

    3: ("Tóm tắt. ",
        "Bài báo này nghiên cứu bài toán tối ưu hóa thiết kế đa mục tiêu cho khung thép "
        "tiết diện chữ I hàn, tuân thủ các tiêu chuẩn TCVN 2737:2023 và TCVN 5575:2024 "
        "hiện hành. Một chương trình Phần tử hữu hạn (FEM) tự xây dựng được kết hợp với "
        "hàm phạt phi tuyến, và ba thuật toán metaheuristic — NSGA-II, PESA-II và MOPSO — "
        "được áp dụng để giải quyết hai mục tiêu đối nghịch: giảm thiểu khối lượng thép "
        "kết cấu và giảm thiểu chuyển vị ngang của khung. Các biên Pareto thu được cho "
        "thấy sự khác biệt rõ rệt trong khả năng khảo sát không gian tìm kiếm của từng "
        "thuật toán; việc xác định “điểm knee-point” trên mỗi biên giúp lựa chọn các "
        "phương án thiết kế khả thi để thi công. Qua 30 lần chạy độc lập cho mỗi thuật "
        "toán, MOPSO hội tụ nhanh hơn 2,3–3,7 lần so với hai phương pháp còn lại, trong "
        "khi PESA-II cho chất lượng biên Pareto tổng thể tốt nhất (hypervolume cao nhất) "
        "và cho nghiệm ở vùng cứng nhất ổn định nhất. Dù vậy, cả ba thuật toán đều hội tụ "
        "về một điểm thỏa hiệp (knee-point) gần như trùng nhau, khoảng 3520 kg và chuyển "
        "vị 6,8 mm, cho thấy đây là một lựa chọn thiết kế bền vững, không phụ thuộc vào "
        "thuật toán sử dụng.", True),

    4: (None, "Từ khóa: Tối ưu hóa đa mục tiêu; khung thép; biên Pareto; TCVN 5575:2024; "
              "metaheuristic; MOPSO; điểm knee-point", False),

    5: (None, "Giới thiệu", False),

    6: (None,
        "Thiết kế khung thép về bản chất là bài toán đánh đổi giữa tính kinh tế vật liệu "
        "và độ cứng ngang. Các phương pháp tối ưu đơn mục tiêu truyền thống thường cho ra "
        "kết cấu quá mềm dẻo, đó là lý do Tối ưu hóa đa mục tiêu (MOO) dựa trên "
        "metaheuristic trở thành hướng tiếp cận được ưa chuộng: nó tạo ra một tập nghiệm "
        "thỏa hiệp Pareto liên tục thay vì một nghiệm cứng nhắc duy nhất. Trong thực tiễn "
        "kỹ thuật tại Việt Nam, tiêu chuẩn TCVN 5575:2024 [1] mới ban hành bổ sung các "
        "ràng buộc phi tuyến nghiêm ngặt về độ mảnh cấu kiện, mất ổn định cục bộ bản bụng/"
        "bản cánh và ổn định ngoài mặt phẳng, khiến không gian thiết kế khả thi trở nên "
        "phức tạp và không liên tục. Để làm việc trong các ràng buộc quy chuẩn này, "
        "nghiên cứu xây dựng một khung tính toán tự động, kết hợp bộ giải Phần tử hữu hạn "
        "(FEM) tự phát triển với ba thuật toán tối ưu đa mục tiêu đã được kiểm chứng: "
        "NSGA-II [3], MOPSO [4] và PESA-II [5]. Các biến thể trí tuệ bầy đàn mới thường "
        "hội tụ nhanh hơn trên các hàm toán học không ràng buộc, nhưng độ tin cậy của "
        "chúng dưới các ràng buộc kết cấu quy chuẩn hóa chặt chẽ chưa được kiểm chứng "
        "nghiêm ngặt; vì vậy nhóm nghiên cứu chủ động lựa chọn ba thuật toán đã được tài "
        "liệu hóa đầy đủ này để đảm bảo độ chính xác toán học của các biên Pareto thu "
        "được. Ngoài việc xác định thuật toán hiệu quả nhất cho hình học tiết diện chữ I "
        "theo tiêu chuẩn mới, nghiên cứu so sánh này còn cung cấp một cơ sở đối chiếu "
        "đáng tin cậy để đánh giá các phát triển thuật toán trong tương lai cho bài toán "
        "tối ưu kết cấu [6].", True),

    7: (None, "Khung tính toán và phương pháp", False),
    8: (None, "Công thức toán học của bài toán kết cấu", False),

    9: (None,
        "Khối lượng khung và chuyển vị ngang được tối ưu hóa đồng thời, sử dụng một "
        "vectơ liên tục tám chiều xác định tiết diện các cấu kiện làm biến thiết kế. Các "
        "yêu cầu về độ bền, ổn định tổng thể và ổn định cục bộ theo TCVN 5575:2024 được "
        "áp đặt thông qua một hàm phạt tĩnh phi tuyến với hệ số nhân lớn (R) nhằm phạt "
        "mọi vi phạm ràng buộc. Bộ giải FEM tự xây dựng tính toán phản ứng kết cấu cho "
        "mỗi phương án thiết kế và phản hồi lại để dẫn dắt ba thuật toán metaheuristic — "
        "NSGA-II, MOPSO và PESA-II — mỗi thuật toán được chạy 30 lần độc lập với cùng "
        "kích thước quần thể 100 cá thể (Mục 2.2).", True),

    12: (None,
         "Ràng buộc biên: 0.20–0.50 m, 0.30–0.80 m, 6–18 mm, 8–22 mm cho tiết diện cột "
         "(x1–x4), và 0.15–0.40 m, 0.25–0.60 m, 5–12 mm, 6–16 mm cho tiết diện dầm "
         "(x5–x8).", False),

    16: (None,
         "Trong thực tế, mười hai điều kiện kiểm tra gj(x) (Mục 2.1) được quy về một tỷ "
         "số bất lợi nhất duy nhất r(x) = maxj gj(x); khi đó, phương trình (4) được áp "
         "dụng dưới dạng nhân tương đương Fk(x) = fk(x)·[1+λ(max(0, r(x)−1))²], với "
         "λ = 5000.", False),

    17: (None, "Tích hợp và cấu hình các thuật toán metaheuristic", False),

    18: (None,
         "Chương trình mô phỏng số được xây dựng dựa trên một bộ giải Phần tử hữu hạn "
         "(FEM) tự phát triển, kết hợp với mô-đun kiểm tra tuân thủ TCVN 5575:2024 viết "
         "trên nền MATLAB, được xây dựng dựa trên các mẫu tính toán tiến hóa Yarpiz [7]. "
         "Mỗi vectơ thiết kế được chuyển đổi thành nội lực và chuyển vị của cấu kiện, và "
         "cơ chế phạt trong phương trình (4) sử dụng phản hồi này để dẫn dắt quá trình "
         "tìm kiếm. NSGA-II và PESA-II tạo ra toàn bộ quần thể con cháu ở mỗi thế hệ, "
         "trong khi MOPSO chỉ cập nhật lại bầy đàn hiện có; do đó, cùng một số vòng lặp "
         "không tương ứng với cùng một khối lượng tính toán ở cả ba phương pháp. Để đảm "
         "bảo tính công bằng khi so sánh, cả ba thuật toán đều sử dụng cùng kích thước "
         "quần thể (100 cá thể) và được hiệu chỉnh trước để tiêu tốn cùng một ngân sách "
         "30.000 lần đánh giá hàm mục tiêu (NFE), tương ứng với 227, 299 và 300 thế hệ "
         "lần lượt cho MOPSO, NSGA-II và PESA-II (Bảng 1). Mỗi thuật toán sau đó được "
         "chạy 30 lần độc lập trong cùng ngân sách này để có thể so sánh thống kê các "
         "tập Pareto và thời gian tính toán thu được.", True),

    19: ("Bảng 1.",
         " Các biến tính toán cụ thể và cấu hình khởi tạo cho các thuật toán NSGA-II, "
         "MOPSO và PESA-II.", False),

    20: (None, "Kết quả thực nghiệm và thảo luận", False),
    21: (None, "Mô tả mô hình và kiểm chứng FEM", False),

    22: (None,
         "Để đánh giá ba khung tính toán tối ưu, nghiên cứu phân tích một khung phẳng "
         "thép hai tầng, một nhịp làm mô hình chuẩn, chịu tác động của tải trọng đứng và "
         "áp lực gió tương đương của vùng. Trước khi đưa mô hình kết cấu này vào vòng "
         "lặp tối ưu hóa, chương trình FEM viết trên nền MATLAB được kiểm chứng trước "
         "với phần mềm thương mại SAP2000. Trên toàn bộ các nút và phần tử được theo "
         "dõi, sai số tương đối lớn nhất về chuyển vị nút chỉ là 0.0445%, cho thấy bộ "
         "giải tự xây dựng đủ chính xác để sử dụng đánh giá phản ứng kết cấu và giá trị "
         "hàm mục tiêu trong suốt quá trình tối ưu hóa.", True),

    23: ("Bảng 2.",
         " Tổng hợp các biến thiết kế và ràng buộc đầu vào cho mô hình khung số.", False),

    25: ("Hình 1.", " Sơ đồ cấu hình khung và tải trọng [2]", False),

    26: (None, "Khảo sát biên Pareto và xác định điểm knee-point", False),

    27: (None,
         "Gộp chung các nghiệm không bị trội của cả ba thuật toán qua 30 lần chạy tạo "
         "thành một biên Pareto lồi, rõ ràng, trải rộng khoảng 2.0–7.5 tấn khối lượng và "
         "3–18 mm chuyển vị (Hình 2). Ở các vùng biên, ba phương pháp chỉ khác biệt "
         "không đáng kể. Về khối lượng nhỏ nhất, PESA-II và MOPSO không khác biệt có ý "
         "nghĩa thống kê (giá trị tốt nhất trong 30 lần chạy lần lượt là 1983.2 kg và "
         "1983.7 kg; trung bình±độ lệch chuẩn là 2074±58 kg và 2108±79 kg), và cả hai "
         "đều nhẹ hơn NSGA-II (tốt nhất 2041 kg, trung bình 2141±70 kg). Về chuyển vị "
         "nhỏ nhất, PESA-II vừa cho kết quả tốt nhất (2.89 mm ứng với 7546 kg) vừa ổn "
         "định nhất (2.99±0.06 mm), vượt trội hơn NSGA-II (3.08 mm) và MOPSO (3.14 mm); "
         "điều này phù hợp với chỉ số hypervolume trung bình cao hơn của PESA-II "
         "(77.25±0.19, so với 76.15±0.69 của NSGA-II và 74.49±1.27 của MOPSO). Ngược "
         "lại, NSGA-II cho biên phân bố đều nhất (chỉ số spacing 0.0068±0.0006). Dù có "
         "những khác biệt này, cả ba thuật toán đều hội tụ về gần như cùng một điểm "
         "knee-point, trong khoảng 3503–3535 kg và 6.75–6.91 mm, cho thấy nghiệm thỏa "
         "hiệp này không phụ thuộc vào thuật toán được sử dụng.", True),

    29: ("Hình 2.",
         " Tập hợp các nghiệm không bị trội (30 lần chạy/thuật toán) và các điểm "
         "knee-point thống nhất trong không gian đánh đổi khối lượng–chuyển vị", False),

    30: (None, "Đánh giá hiệu quả tính toán", False),

    31: (None,
         "Việc gọi lặp lại bộ giải phần tử hữu hạn bên trong vòng lặp tối ưu hóa đòi hỏi "
         "chi phí tính toán lớn. Đo thời gian trên cùng 30 lần chạy (Intel Core "
         "i5-1145G7 @ 2.60 GHz, MATLAB R2023b) cho thấy MOPSO hiệu quả vượt trội: thời "
         "gian chạy trung bình 32.95 giây, nhanh hơn 3.67 lần so với NSGA-II (120.82 "
         "giây) và nhanh hơn 2.29 lần so với PESA-II (75.49 giây). Ngay cả lần chạy "
         "chậm nhất của MOPSO (49.4 giây) cũng nhanh hơn lần chạy nhanh nhất của "
         "NSGA-II (69.2 giây), cho thấy đây là một lợi thế tốc độ ổn định chứ không "
         "phải ngẫu nhiên. Mặc dù PESA-II cho chất lượng nghiệm tổng thể tốt hơn đôi "
         "chút (Mục 3.2), lợi thế về tốc độ giúp MOPSO trở thành lựa chọn thực tiễn hơn "
         "cho các bài toán tối ưu lặp lại nhiều lần, chịu ràng buộc quy chuẩn.", True),

    32: ("Bảng 3.",
         " Các chỉ số thống kê về thời gian tính toán (giây) qua 30 lần chạy độc lập.",
         False),

    33: (None, "Kết luận", False),

    34: (None,
         "Nghiên cứu này trình bày một khung tính toán tối ưu hóa tự động dựa trên FEM "
         "cho khung thép tiết diện chữ I được thiết kế theo TCVN 5575:2024. Các kết "
         "luận chính như sau.", True),

    35: (None,
         "Hiệu năng thuật toán. MOPSO nhanh hơn rõ rệt (2.3–3.7 lần, Mục 3.3), trong "
         "khi PESA-II đạt độ phủ biên Pareto tốt nhất (hypervolume cao nhất) và cho "
         "nghiệm ở vùng cứng nhất ổn định nhất; MOPSO và PESA-II đạt khối lượng nhỏ "
         "nhất gần như nhau (khoảng 1983 kg). Không có thuật toán nào vượt trội ở tất "
         "cả các tiêu chí.", True),

    36: (None,
         "Ứng dụng thiết kế thực tiễn. Cả ba thuật toán đều hội tụ về một điểm "
         "knee-point chung, khoảng 3503–3535 kg với chuyển vị ngang 6.75–6.91 mm, thể "
         "hiện một phương án thiết kế thỏa hiệp bền vững, không phụ thuộc vào thuật "
         "toán được lựa chọn.", True),

    37: (None,
         "Phương pháp luận đề xuất có thể mở rộng cho các kết cấu không gian phức tạp "
         "hơn. Nghiên cứu tiếp theo sẽ bổ sung hiệu ứng độ cứng hình học (P–Δ) và Tối "
         "ưu hóa thiết kế theo độ tin cậy (RBDO) để xét đến tính bất định của đặc trưng "
         "vật liệu và tải trọng tác dụng.", True),

    38: (None, "Tài liệu tham khảo", False),
}

# Cac doan van co cong thuc <m:oMath> chen giua cau: chi thay cac doan
# van ban xen giua, GIU NGUYEN cac phan tu oMath (x1..x8, rho, Ai, Li,
# Ne, Fk, fk, gj) o dung vi tri cu.
SEGMENTS = {
    11: [
        "Trong đó ",
        " lần lượt biểu thị bề rộng bản cánh, chiều cao bản bụng, chiều dày bản bụng và "
        "chiều dày bản cánh của tiết diện cột, còn ",
        " biểu thị các thông số hình học tương ứng của tiết diện dầm.",
    ],
    13: [
        "Trong đó ",
        " là khối lượng riêng của vật liệu thép; ",
        " và ",
        " lần lượt là diện tích tiết diện ngang và chiều dài của phần tử thứ i; và ",
        " là tổng số phần tử cấu thành hệ khung.",
    ],
    14: [
        "Trong đó ",
        " là giá trị chuyển vị ngang lớn nhất tại nút đỉnh của khung, được trích xuất "
        "trực tiếp từ bộ giải phần tử hữu hạn (FEM).",
    ],
    15: [
        "Trong đó ",
        " và ",
        " lần lượt là hàm mục tiêu đã phạt và hàm mục tiêu gốc ứng với tiêu chí thứ k; "
        "R là hệ số phạt tĩnh có giá trị lớn; và ",
        " biểu thị ràng buộc ứng xử thứ j được đánh giá trong mô hình.",
    ],
}

# --------------------------------------------------------------------
# 4) Nhan (label) trong bang so lieu can dich - KHONG dong vao cot
#    ky hieu/gia tri (chua dinh dang subscript) va khong dong vao du lieu.
# --------------------------------------------------------------------

TABLE1_LABELS = {  # cot 2 (Algorithm parameters) cua Bang 1 = doc.tables[4]
    "Item": "Mục",
    "Algorithm parameters": "Thông số thuật toán",
    "Population size (Npop​)": "Kích thước quần thể (Npop)",
    "Maximum iterations (Itermax​)": "Số vòng lặp tối đa (Itermax)",
    "Inertia weight (w)": "Trọng số quán tính (w)",
    "Acceleration coefficients (c1​,c2​)": "Hệ số gia tốc (c1, c2)",
    "Crossover rate (pcross​)": "Tỷ lệ lai ghép (pcross)",
    "Mutation rate (pmut​)": "Tỷ lệ đột biến (pmut)",
    "External grid size": "Kích thước lưới ngoài",
}

TABLE2_LABELS = {  # cot 1 va 2 cua Bang 2 = doc.tables[5]
    "No.": "STT",
    "Model parameters": "Thông số mô hình",
    "Main span / Frame spacing": "Nhịp chính / Bước khung",
    "Story 1 height / Story 2 height": "Chiều cao tầng 1 / tầng 2",
    "Total vertical load (dead + live loads)": "Tổng tải trọng đứng (tĩnh tải + hoạt tải)",
    "Basic wind pressure [2]": "Áp lực gió cơ bản [2]",
    "Yield strength (SS400 steel)": "Cường độ chảy (thép SS400)",
    "Elastic modulus (SS400 steel)": "Mô đun đàn hồi (thép SS400)",
}

TABLE3_LABELS = {  # hang tieu de cua Bang 3 = doc.tables[7]
    "Algorithm": "Thuật toán",
    "Mean time": "Thời gian trung bình",  # doan 1/2 cua o tieu de (2 doan van)
    "(tmean​)": "(tmean)",                 # doan 2/2 cua o tieu de
    "Mean time (tmean​)": "Thời gian trung bình (tmean)",  # phong khi la 1 doan
    "Standard deviation (σ)": "Độ lệch chuẩn (σ)",
    "Fastest run (tmin​)": "Lần chạy nhanh nhất (tmin)",
    "Slowest run (tmax​)": "Lần chạy chậm nhất (tmax)",
}

TABLE2_HEADER_LABELS = {  # cot 3,4 cua Bang 2 = doc.tables[5] (chi hang tieu de)
    "Symbol": "Ký hiệu",
    "Applied value": "Giá trị áp dụng",
}


def set_cell_text_if_match(cell, label_map):
    """Doi khop tung DOAN VAN (paragraph) ben trong o, khong phai toan bo
    van ban cua o - vi 1 o co the gom nhieu doan (vi du o tieu de Bang 3
    'Mean time' + '(tmean)' nam tren 2 dong/2 doan rieng)."""
    changed = False
    for p in cell.paragraphs:
        original = p.text
        if original in label_map:
            set_paragraph_text(p, label_map[original])
            changed = True
    return changed


def main():
    document = docx.Document(SRC)
    paragraphs = document.paragraphs

    # --- 4.1 doan van xuoi (khong cong thuc chen giua) ---
    for idx, item in WHOLE.items():
        if len(item) == 3:
            bold_prefix, text, force_10pt = item
        else:
            bold_prefix, text = item
            force_10pt = False
        p = paragraphs[idx]
        set_paragraph_text(p, text, bold_prefix=bold_prefix,
                            font_pt=10 if force_10pt else None)

    # --- 4.2 doan van co cong thuc oMath chen giua ---
    for idx, seg_texts in SEGMENTS.items():
        p_el = paragraphs[idx]._p
        replace_text_segments(p_el, seg_texts)

    # --- 4.3 nhan trong bang so lieu ---
    t1 = document.tables[4]  # Bang 1: thong so thuat toan
    for row in t1.rows:
        set_cell_text_if_match(row.cells[0], TABLE1_LABELS)  # header "Item"
        set_cell_text_if_match(row.cells[1], TABLE1_LABELS)

    t2 = document.tables[5]  # Bang 2: thong so mo hinh
    for row in t2.rows:
        set_cell_text_if_match(row.cells[0], TABLE2_LABELS)
        set_cell_text_if_match(row.cells[1], TABLE2_LABELS)
        set_cell_text_if_match(row.cells[2], TABLE2_HEADER_LABELS)
        set_cell_text_if_match(row.cells[3], TABLE2_HEADER_LABELS)

    t3 = document.tables[7]  # Bang 3: thong ke thoi gian chay
    for row in t3.rows:
        for cell in row.cells:
            set_cell_text_if_match(cell, TABLE3_LABELS)

    document.save(OUT)
    print(f"Da tao: {OUT}")
    print(f"So doan van thay toan bo: {len(WHOLE)}")
    print(f"So doan van thay theo doan (giu cong thuc): {len(SEGMENTS)}")


if __name__ == "__main__":
    main()
