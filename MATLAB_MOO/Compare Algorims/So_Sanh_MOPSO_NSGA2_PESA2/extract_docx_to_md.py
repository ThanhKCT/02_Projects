"""
extract_docx_to_md.py
----------------------
Trich xuat toan bo noi dung file "ICERA. V2.docx" sang Markdown ("ICERA. V3.md")
bang thu vien python-docx, giu nguyen:
  - Tieu de (heading1/heading2/papertitle -> #, ##, ###)
  - Bang bieu (chuyen thanh bang Markdown)
  - Cong thuc toan (OMML -> LaTeX, boc trong $...$ hoac $$...$$)
  - Hinh anh (trich xuat ra thu muc media va chen lien ket Markdown)

Chay: python extract_docx_to_md.py
"""

import os
import re
import docx
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

SRC = "ICERA. V2.docx"
OUT_MD = "ICERA. V3.md"
MEDIA_DIR = "ICERA_V3_media"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def local(tag):
    """Tra ve ten the khong co namespace, vi du '{...}oMath' -> 'oMath'."""
    return tag.split("}")[-1] if "}" in tag else tag


# --------------------------------------------------------------------------
# 1) Chuyen doi cong thuc toan (OMML) sang chuoi LaTeX gan dung
# --------------------------------------------------------------------------

def omml_children_to_latex(el):
    if el is None:
        return ""
    return "".join(omml_to_latex(c) for c in el)


def _find(el, name):
    return el.find(qn(f"m:{name}"))


def _findall(el, name):
    return el.findall(qn(f"m:{name}"))


def omml_to_latex(el):
    tag = local(el.tag)

    if tag in ("oMath", "oMathPara"):
        return "".join(omml_to_latex(c) for c in el if local(c.tag) != "oMathParaPr")

    if tag == "r":
        # Van ban trong 1 "run" cong thuc: lay het cac the <m:t>
        return "".join(t.text or "" for t in el if local(t.tag) == "t")

    if tag == "f":
        num = omml_children_to_latex(_find(el, "num"))
        den = omml_children_to_latex(_find(el, "den"))
        return f"\\frac{{{num}}}{{{den}}}"

    if tag == "sSub":
        e = omml_children_to_latex(_find(el, "e"))
        sub = omml_children_to_latex(_find(el, "sub"))
        return f"{{{e}}}_{{{sub}}}"

    if tag == "sSup":
        e = omml_children_to_latex(_find(el, "e"))
        sup = omml_children_to_latex(_find(el, "sup"))
        return f"{{{e}}}^{{{sup}}}"

    if tag == "sSubSup":
        e = omml_children_to_latex(_find(el, "e"))
        sub = omml_children_to_latex(_find(el, "sub"))
        sup = omml_children_to_latex(_find(el, "sup"))
        return f"{{{e}}}_{{{sub}}}^{{{sup}}}"

    if tag == "rad":
        e = omml_children_to_latex(_find(el, "e"))
        deg_el = _find(el, "deg")
        rad_pr = _find(el, "radPr")
        hide_deg = False
        if rad_pr is not None:
            deg_hide = _find(rad_pr, "degHide")
            if deg_hide is not None:
                val = deg_hide.get(qn("m:val"))
                hide_deg = val is None or val in ("1", "true", "on")
        if deg_el is not None and not hide_deg:
            deg = omml_children_to_latex(deg_el)
            if deg.strip():
                return f"\\sqrt[{deg}]{{{e}}}"
        return f"\\sqrt{{{e}}}"

    if tag == "nary":
        pr = _find(el, "naryPr")
        op = "\\sum"
        if pr is not None:
            chr_el = _find(pr, "chr")
            if chr_el is not None:
                sym = chr_el.get(qn("m:val")) or ""
                op = {
                    "∑": "\\sum",
                    "∏": "\\prod",
                    "∫": "\\int",
                    "∬": "\\iint",
                    "∭": "\\iiint",
                    "⋃": "\\bigcup",
                    "⋂": "\\bigcap",
                }.get(sym, sym or "\\sum")
        sub = omml_children_to_latex(_find(el, "sub"))
        sup = omml_children_to_latex(_find(el, "sup"))
        e = omml_children_to_latex(_find(el, "e"))
        out = op
        if sub.strip():
            out += f"_{{{sub}}}"
        if sup.strip():
            out += f"^{{{sup}}}"
        return f"{out} {e}"

    if tag == "d":
        pr = _find(el, "dPr")
        beg, end = "(", ")"
        if pr is not None:
            b = _find(pr, "begChr")
            e_ = _find(pr, "endChr")
            if b is not None and b.get(qn("m:val")) is not None:
                beg = b.get(qn("m:val"))
            if e_ is not None and e_.get(qn("m:val")) is not None:
                end = e_.get(qn("m:val"))
        parts = [omml_children_to_latex(e) for e in _findall(el, "e")]
        return f"{beg}{', '.join(parts)}{end}"

    if tag == "func":
        fname = omml_children_to_latex(_find(el, "fName"))
        e = omml_children_to_latex(_find(el, "e"))
        return f"{fname}({e})"

    if tag in ("limLow", "limUpp"):
        e = omml_children_to_latex(_find(el, "e"))
        lim = omml_children_to_latex(_find(el, "lim"))
        sep = "_" if tag == "limLow" else "^"
        return f"{e}{sep}{{{lim}}}"

    if tag == "bar":
        e = omml_children_to_latex(_find(el, "e"))
        return f"\\overline{{{e}}}"

    if tag == "acc":
        e = omml_children_to_latex(_find(el, "e"))
        return f"\\hat{{{e}}}"

    if tag == "groupChr":
        return omml_children_to_latex(_find(el, "e"))

    if tag == "eqArr":
        rows = [omml_children_to_latex(e) for e in _findall(el, "e")]
        return " \\\\ ".join(rows)

    if tag == "m":  # ma tran
        rows = []
        for mr in _findall(el, "mr"):
            cells = [omml_children_to_latex(e) for e in _findall(mr, "e")]
            rows.append(" & ".join(cells))
        return "\\begin{matrix}" + " \\\\ ".join(rows) + "\\end{matrix}"

    # Mac dinh: de quy vao cac phan tu con (bao gom m:e don le, box, phant,...)
    return omml_children_to_latex(el)


def render_omath_element(el, display):
    latex = omml_to_latex(el).strip()
    if not latex:
        return ""
    return f"$${latex}$$" if display else f"${latex}$"


# --------------------------------------------------------------------------
# 2) Trich xuat hinh anh
# --------------------------------------------------------------------------

_image_cache = {}
_image_counter = [0]


def extract_image(document, r_embed_id, out_dir):
    if not r_embed_id:
        return None
    if r_embed_id in _image_cache:
        return _image_cache[r_embed_id]
    try:
        part = document.part.related_parts[r_embed_id]
    except KeyError:
        return None
    _image_counter[0] += 1
    ext = os.path.splitext(part.partname)[1] or ".bin"
    filename = f"image{_image_counter[0]}{ext}"
    os.makedirs(out_dir, exist_ok=True)
    with open(os.path.join(out_dir, filename), "wb") as f:
        f.write(part.blob)
    rel_path = f"{os.path.basename(out_dir)}/{filename}"
    _image_cache[r_embed_id] = rel_path
    return rel_path


def render_drawing_or_pict(document, el, out_dir):
    """el la the <w:drawing> hoac <w:pict> - tim r:embed va chen anh Markdown."""
    r_attr = qn("r:embed")
    links = []
    for blip in el.iter():
        if local(blip.tag) in ("blip", "imagedata"):
            rid = blip.get(r_attr) or blip.get(qn("r:id"))
            path = extract_image(document, rid, out_dir)
            if path:
                links.append(f"![]({path})")
    return " ".join(links)


# --------------------------------------------------------------------------
# 3) Render mot doan van (paragraph) sang Markdown, giu dung thu tu
#    van ban <-> cong thuc <-> hinh anh
# --------------------------------------------------------------------------

def _wrap_marker(text, marker):
    """Boc marker (**, _) sat noi dung, dua khoang trang/newline ra ngoai
    de dung cu phap Markdown (delimiter khong duoc dung ke khoang trang)."""
    core = text.strip()
    if not core:
        return text
    lead_len = len(text) - len(text.lstrip())
    trail_len = len(text) - len(text.rstrip())
    leading = text[:lead_len]
    trailing = text[len(text) - trail_len:] if trail_len else ""
    return f"{leading}{marker}{core}{marker}{trailing}"


def render_run(r_el, document=None, media_dir=None):
    """Ghep noi dung text cua 1 <w:r>, ap dung dam/nghieng don gian.
    Mot <w:r> co the chua <w:drawing>/<w:pict> (hinh anh) long ben trong."""
    text_parts = []
    image_parts = []
    for child in r_el:
        t = local(child.tag)
        if t == "t":
            text_parts.append(child.text or "")
        elif t == "tab":
            text_parts.append("\t")
        elif t in ("br", "cr"):
            text_parts.append("  \n")
        elif t == "noBreakHyphen":
            text_parts.append("-")
        elif t in ("drawing", "pict") and document is not None:
            img_md = render_drawing_or_pict(document, child, media_dir)
            if img_md:
                image_parts.append(img_md)
    text = "".join(text_parts)
    if not text:
        return " ".join(image_parts)

    bold = italic = False
    rpr = r_el.find(qn("w:rPr"))
    if rpr is not None:
        if rpr.find(qn("w:b")) is not None and rpr.find(qn("w:b")).get(qn("w:val")) != "0":
            bold = True
        if rpr.find(qn("w:i")) is not None and rpr.find(qn("w:i")).get(qn("w:val")) != "0":
            italic = True

    if text.strip():
        if bold:
            text = _wrap_marker(text, "**")
        if italic:
            text = _wrap_marker(text, "_")
    if image_parts:
        text = (text + " " + " ".join(image_parts)).strip()
    return text


def render_paragraph_inline(p_el, document, media_dir):
    """Duyet cac phan tu con truc tiep cua <w:p> theo dung thu tu xuat hien."""
    parts = []
    for child in p_el:
        tag = local(child.tag)
        if tag == "r":
            parts.append(render_run(child, document, media_dir))
        elif tag == "hyperlink":
            for sub in child:
                if local(sub.tag) == "r":
                    parts.append(render_run(sub, document, media_dir))
        elif tag == "oMath":
            parts.append(render_omath_element(child, display=False))
        elif tag == "oMathPara":
            for sub in child:
                if local(sub.tag) == "oMath":
                    parts.append(render_omath_element(sub, display=True))
        elif tag in ("drawing", "pict"):
            parts.append(render_drawing_or_pict(document, child, media_dir))
        # bo qua bookmarkStart/End, proofErr, commentReference, ...
    return "".join(parts).strip()


# --------------------------------------------------------------------------
# 4) Xac dinh cap tieu de tu ten style
# --------------------------------------------------------------------------

def heading_level(style_name):
    if not style_name:
        return None
    name = style_name.strip().lower()
    if name in ("papertitle", "title", "doc title"):
        return 1
    m = re.match(r"heading\s*([0-9]+)", name)
    if m:
        return int(m.group(1)) + 1  # heading1 -> ##, heading2 -> ###, ...
    return None


# --------------------------------------------------------------------------
# 5) Duyet cac khoi (paragraph / table) theo dung thu tu trong body
# --------------------------------------------------------------------------

def iter_block_items(document):
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


EQ_NUM_RE = re.compile(r"^\(\d+\)$")


def table_is_equation_wrapper(table):
    """Phat hien bang duoc dung nhu 'khung' chua cong thuc + so thu tu,
    kieu trinh bay pho bien trong template IEEE/ICERA."""
    if len(table.rows) != 1:
        return False
    xml = table._tbl.xml
    if "oMath" not in xml:
        return False
    last_cell_text = table.rows[-1].cells[-1].text.strip()
    return bool(EQ_NUM_RE.match(last_cell_text))


def render_table_as_equation(table, document, media_dir):
    row = table.rows[0]
    eq_num = row.cells[-1].text.strip()
    formula_parts = []
    for cell in row.cells[:-1]:
        for p in cell.paragraphs:
            rendered = render_paragraph_inline(p._p, document, media_dir)
            if rendered:
                formula_parts.append(rendered)
    formula = " ".join(formula_parts).strip()
    if not formula:
        return ""
    # Neu formula chua duoc boc $..$ (vi du cell rong nhung co oMath rieng), boc lai
    if not (formula.startswith("$")):
        formula = f"${formula}$"
    return f"{formula}  {eq_num}"


def escape_md_cell(text):
    text = text.replace("\n", "<br>").replace("|", "\\|")
    return text.strip()


def render_table_as_markdown(table, document, media_dir):
    rows_out = []
    for row in table.rows:
        cells_out = []
        for cell in row.cells:
            cell_parts = []
            for p in cell.paragraphs:
                rendered = render_paragraph_inline(p._p, document, media_dir)
                if rendered:
                    cell_parts.append(rendered)
            cells_out.append(escape_md_cell(" ".join(cell_parts)))
        rows_out.append(cells_out)

    if not rows_out:
        return ""

    ncols = max(len(r) for r in rows_out)
    for r in rows_out:
        while len(r) < ncols:
            r.append("")

    lines = []
    header = rows_out[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * ncols) + " |")
    for r in rows_out[1:]:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


# --------------------------------------------------------------------------
# 6) Chuong trinh chinh
# --------------------------------------------------------------------------

def main():
    document = docx.Document(SRC)
    md_lines = []

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            style_name = block.style.name if block.style is not None else None
            level = heading_level(style_name)
            text = render_paragraph_inline(block._p, document, MEDIA_DIR)
            if not text:
                continue
            if level:
                md_lines.append(f"{'#' * level} {text}")
            else:
                md_lines.append(text)
            md_lines.append("")

        elif isinstance(block, Table):
            if table_is_equation_wrapper(block):
                eq_md = render_table_as_equation(block, document, MEDIA_DIR)
                if eq_md:
                    md_lines.append(eq_md)
                    md_lines.append("")
            else:
                table_md = render_table_as_markdown(block, document, MEDIA_DIR)
                if table_md:
                    md_lines.append(table_md)
                    md_lines.append("")

    content = "\n".join(md_lines).strip() + "\n"
    with open(OUT_MD, "w", encoding="utf-8") as f:
        f.write(content)

    print(f"Da tao xong: {OUT_MD}")
    print(f"So dong Markdown: {len(md_lines)}")
    if os.path.isdir(MEDIA_DIR):
        print(f"Hinh anh trich xuat vao: {MEDIA_DIR}/ ({len(os.listdir(MEDIA_DIR))} file)")


if __name__ == "__main__":
    main()
