# -*- coding: utf-8 -*-
"""
revise_wording_v4.py
---------------------
Doc "ICERA. V2.docx" bang python-docx, thay lai VAN PHONG (khong doi bat ky
so lieu/ket qua/trich dan nao) cho cac doan van xuoi (abstract, gioi thieu,
thao luan, ket luan), giu nguyen 100% tieu de, bang bieu, cong thuc, chu
thich hinh/bang va danh muc tai lieu tham khao. Xuat ra "ICERA. V4.docx".

Tat ca so lieu trong ban thay the da duoc doi chieu voi du lieu goc
(BangKetQua_30Lan.xlsx, KetQuaThongKe_ThuatToan.xlsx,
BaoCao_DoiChieu_TongHop. SAP.xlsx) truoc khi ap dung - xem bao cao ra soat.
"""

import docx

SRC = "ICERA. V2.docx"
OUT = "ICERA. V4.docx"

# { paragraph_index_trong_document : van_ban_moi }
# Cac chi so lay tu doc._element.body (giong docx.Document(...).paragraphs)
REPLACEMENTS = {
    3: (
        "Abstract. ",
        "This paper investigates the multi-objective design optimization of welded "
        "I-section steel frames under the current TCVN 2737:2023 and TCVN 5575:2024 "
        "standards. A custom Finite Element Method (FEM) solver is coupled with a "
        "non-linear penalty function, and three metaheuristic algorithms — NSGA-II, "
        "PESA-II, and MOPSO — are applied to two conflicting objectives: minimizing "
        "structural steel weight and minimizing lateral frame displacement. The "
        "resulting Pareto fronts reveal clear differences in how each algorithm "
        "explores the search space; identifying the “knee-point” of each front "
        "helps select design configurations that are practical to build. Over 30 "
        "independent runs per algorithm, MOPSO converges 2.3–3.7 times faster than "
        "the other two methods, while PESA-II produces the best overall Pareto-front "
        "quality (highest hypervolume) and the most repeatable extreme-stiffness "
        "solutions. All three algorithms nevertheless settle on a nearly identical "
        "knee-point compromise, around 3520 kg and 6.8 mm of drift, which points to "
        "a robust, algorithm-independent design choice for practical steel-frame "
        "design."
    ),
    6: (
        None,
        "Designing a steel frame is essentially a trade-off between material economy "
        "and lateral stiffness. Traditional single-objective methods tend to produce "
        "overly flexible structures, which is why Multi-Objective Optimization (MOO) "
        "based on metaheuristics has become the preferred approach: it yields a "
        "continuous Pareto trade-off instead of a single rigid solution. For "
        "engineering practice in Vietnam, the newly enacted TCVN 5575:2024 [1] adds "
        "severe non-linear constraints on member slenderness, local plate buckling, "
        "and out-of-plane stability, turning the feasible design space into a highly "
        "complex, discontinuous domain. To work within these code boundaries, this "
        "study builds an automated computational framework that couples an in-house "
        "Finite Element Method (FEM) engine with three well-established "
        "multi-objective solvers: NSGA-II [3], MOPSO [4], and PESA-II [5]. Newer "
        "swarm-intelligence variants often converge faster on unconstrained "
        "mathematical benchmarks, but their reliability under heavily codified "
        "structural constraints has not been rigorously verified; for this reason "
        "we deliberately choose these three well-documented algorithms, so that the "
        "mathematical accuracy of the resulting Pareto fronts can be trusted. Beyond "
        "identifying the most efficient solver for I-section geometries under the "
        "new code, this comparison also provides a reliable baseline against which "
        "future algorithmic developments in structural optimization can be "
        "evaluated [6]."
    ),
    9: (
        None,
        "Frame mass and lateral drift are optimized simultaneously, using an "
        "eight-dimensional continuous vector that defines the member cross-sections "
        "as design variables. Strength, global stability, and local buckling "
        "requirements from TCVN 5575:2024 are enforced through a non-linear static "
        "penalty function with a large multiplier (R) that penalizes any constraint "
        "violation. The in-house FEM solver evaluates the structural response for "
        "each candidate design and feeds it back to guide the three metaheuristics "
        "— NSGA-II, MOPSO, and PESA-II — each run for 30 independent trials with "
        "the same population size of 100 individuals (Section 2.2)."
    ),
    18: (
        None,
        "The numerical simulation is built around an in-house Finite Element Method "
        "(FEM) code coupled with a MATLAB-based TCVN 5575:2024 compliance-checking "
        "module, implemented on top of the Yarpiz evolutionary-computation "
        "templates [7]. Each design vector is translated into member forces and "
        "displacements, and the penalty mechanism of Eq. (4) uses this response to "
        "steer the search. NSGA-II and PESA-II generate a full offspring population "
        "every generation, whereas MOPSO only updates its existing swarm; "
        "consequently, the same iteration count does not represent the same "
        "computational effort for all three methods. To keep the comparison fair, "
        "all three algorithms use the same population size (100 individuals) and "
        "are each calibrated beforehand to spend the same budget of 30,000 "
        "objective-function evaluations (NFE), which corresponds to 227, 299, and "
        "300 generations for MOPSO, NSGA-II, and PESA-II, respectively (Table 1). "
        "Each algorithm is then run for 30 independent trials under this shared "
        "budget so that the resulting Pareto sets and runtimes can be compared "
        "statistically."
    ),
    22: (
        None,
        "To evaluate the three optimization frameworks, we analyze a benchmark "
        "two-story, single-span steel plane frame subjected to vertical gravity "
        "loads and the regional equivalent wind pressure. Before this structural "
        "model is embedded in the optimization loops, the in-house MATLAB-based "
        "FEM code is first checked against the commercial software SAP2000. Across "
        "all monitored nodes and elements, the maximum relative error in nodal "
        "displacement is only 0.0445%, which indicates that the in-house solver is "
        "precise enough to be used for evaluating structural response and "
        "objective values throughout the optimization runs."
    ),
    27: (
        None,
        "Pooling the non-dominated solutions from all three algorithms over the 30 "
        "runs produces a well-defined, convex Pareto front spanning roughly "
        "2.0–7.5 t of mass and 3–18 mm of drift (Fig. 2). At the extremes, the "
        "three methods behave only slightly differently. For minimum mass, "
        "PESA-II and MOPSO are statistically indistinguishable (best-of-30 values "
        "of 1983.2 kg and 1983.7 kg; mean±SD of 2074±58 kg and 2108±79 kg), and "
        "both are lighter than NSGA-II (2041 kg best, 2141±70 kg mean). For "
        "minimum displacement, PESA-II is both the best (2.89 mm at 7546 kg) and "
        "the most repeatable (2.99±0.06 mm), ahead of NSGA-II (3.08 mm) and MOPSO "
        "(3.14 mm); this is consistent with PESA-II's higher average hypervolume "
        "(77.25±0.19, versus 76.15±0.69 for NSGA-II and 74.49±1.27 for MOPSO). "
        "NSGA-II, on the other hand, produces the most evenly spaced front "
        "(spacing of 0.0068±0.0006). Even with these differences, all three "
        "algorithms converge on almost the same knee-point, between 3503 and 3535 "
        "kg and between 6.75 and 6.91 mm, suggesting that this compromise solution "
        "does not depend on which algorithm is used."
    ),
    31: (
        None,
        "Repeatedly calling the finite element solver inside the optimization loop "
        "is computationally expensive. Timing the same 30 trials (Intel Core "
        "i5-1145G7 @ 2.60 GHz, MATLAB R2023b) shows that MOPSO is markedly more "
        "efficient: its average runtime of 32.95 s is 3.67 times faster than "
        "NSGA-II (120.82 s) and 2.29 times faster than PESA-II (75.49 s). Even "
        "MOPSO's slowest trial (49.4 s) finishes before NSGA-II's fastest one "
        "(69.2 s), which points to a consistent speed advantage rather than a "
        "one-off result. Although PESA-II gives slightly better solution quality "
        "overall (Section 3.2), MOPSO's speed advantage makes it the more "
        "practical choice for repeated, code-constrained optimization runs."
    ),
    34: (
        None,
        "This study presents an automated, FEM-based optimization framework for "
        "I-section steel frames designed under TCVN 5575:2024. The main "
        "conclusions are as follows."
    ),
    35: (
        None,
        "Solver performance. MOPSO is decisively faster (2.3–3.7×, Section 3.3), "
        "while PESA-II achieves the best Pareto-front coverage (highest "
        "hypervolume) and the most consistent stiffest-design solutions; MOPSO and "
        "PESA-II reach almost the same minimum mass (about 1983 kg). None of the "
        "three algorithms dominates on every criterion."
    ),
    36: (
        None,
        "Practical design application. All three algorithms converge on a common "
        "knee-point, 3503–3535 kg with 6.75–6.91 mm of lateral drift, which "
        "represents a robust design compromise that does not depend on the chosen "
        "algorithm."
    ),
    37: (
        None,
        "The proposed methodology can be extended to more complex spatial "
        "structures. Future work will incorporate geometric-stiffness (P–Δ) "
        "effects and Reliability-Based Design Optimization (RBDO) to account for "
        "uncertainty in material properties and applied loads."
    ),
}


def set_paragraph_text(paragraph, new_text, bold_prefix=None, base_font_pt=10):
    """Xoa toan bo run hien co va thay bang noi dung moi, giu font size goc.
    Neu co bold_prefix (vi du 'Abstract. ') thi tao rieng 1 run in dam."""
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    if bold_prefix:
        r1 = paragraph.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = docx.shared.Pt(base_font_pt)
        r2 = paragraph.add_run(new_text)
        r2.font.size = docx.shared.Pt(base_font_pt)
    else:
        r = paragraph.add_run(new_text)
        r.font.size = docx.shared.Pt(base_font_pt)


def main():
    document = docx.Document(SRC)
    paragraphs = document.paragraphs

    applied = []
    for idx, (bold_prefix, new_text) in REPLACEMENTS.items():
        p = paragraphs[idx]
        old_preview = p.text.strip()[:60]
        set_paragraph_text(p, new_text, bold_prefix=bold_prefix)
        applied.append((idx, old_preview))

    document.save(OUT)

    print(f"Da tao: {OUT}")
    print(f"So doan van da chinh sua van phong: {len(applied)}")
    for idx, preview in applied:
        print(f"  - [{idx}] {preview}...")


if __name__ == "__main__":
    main()
